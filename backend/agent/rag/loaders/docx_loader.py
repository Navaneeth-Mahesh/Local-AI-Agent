from pathlib import Path

from docx import Document as DocxDocument

from agent.rag.loaders.base import (
    BaseDocumentLoader,
)
from agent.rag.models import Document


class DocxLoader(
    BaseDocumentLoader,
):

    def supports(
        self,
        path: Path,
    ) -> bool:

        return path.suffix.lower() == ".docx"

    def load(
        self,
        path: Path,
    ) -> Document:

        doc = DocxDocument(path)

        text = "\n".join(
            paragraph.text
            for paragraph in doc.paragraphs
        )

        return Document(
            path=path,
            title=path.stem,
            content=text,
            metadata={},
        )